"""
Sistema LIA - Rotas de Exportação
==================================
Gera arquivos DOCX e visualização de impressão para artefatos.
"""

from fastapi import APIRouter, Depends, HTTPException, Request
from fastapi.responses import StreamingResponse, HTMLResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
import asyncio
from app.database import get_db
from app.models.projeto import Projeto
from app.models.user import User
from app.auth import current_active_user as auth_get_current_user
from app.models.artefatos import ARTEFATO_MAP
from app.routers.views.common import templates
import io
from datetime import datetime

# Import do service de PDF
try:
    from app.services.pdf_service import gerar_pdf_artefato, gerar_nome_arquivo_pdf, HAS_WEASYPRINT
except ImportError:
    HAS_WEASYPRINT = False

# Tenta importar python-docx
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

router = APIRouter()

@router.get("/{projeto_id}/{tipo_artefato}/docx")
async def exportar_docx(
    projeto_id: int,
    tipo_artefato: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(auth_get_current_user)
):
    """Gera arquivo DOCX do artefato"""
    if not HAS_DOCX:
        raise HTTPException(status_code=501, detail="Biblioteca python-docx não instalada no servidor.")

    if tipo_artefato not in ARTEFATO_MAP:
        raise HTTPException(status_code=404, detail="Tipo de artefato inválido")

    config = ARTEFATO_MAP[tipo_artefato]
    Model = config["model"]
    campos_config = config["config"]
    titulo_doc = config["titulo"]

    result = await db.execute(select(Projeto).filter(Projeto.id == projeto_id))
    projeto = result.scalars().first()
    if not projeto:
        raise HTTPException(status_code=404, detail="Projeto não encontrado")

    result = await db.execute(select(Model).filter(Model.projeto_id == projeto_id))
    artefato = result.scalars().first()
    if not artefato:
        raise HTTPException(status_code=404, detail="Artefato não encontrado (salve o rascunho primeiro)")

    # Função auxiliar para gerar DOCX (CPU-bound)
    def _gerar_docx():
        document = Document()
        
        # Título
        heading = document.add_heading(titulo_doc, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadados
        p = document.add_paragraph()
        p.add_run(f"Projeto: ").bold = True
        p.add_run(projeto.titulo)
        
        p = document.add_paragraph()
        p.add_run(f"Versão: ").bold = True
        p.add_run(str(artefato.versao))
        
        p = document.add_paragraph()
        p.add_run(f"Data de Emissão: ").bold = True
        p.add_run(datetime.now().strftime('%d/%m/%Y %H:%M'))
        
        document.add_paragraph("-" * 60)

        # Conteúdo
        for campo, conf in campos_config.items():
            valor = getattr(artefato, campo, "")
            
            document.add_heading(conf.get("label", campo), level=1)
            
            if valor:
                for linha in str(valor).split('\n'):
                    if linha.strip():
                        document.add_paragraph(linha.strip())
            else:
                document.add_paragraph("[Não preenchido]", style="Body Text")

        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        return file_stream

    # Executar em thread pool para não bloquear loop
    file_stream = await asyncio.to_thread(_gerar_docx)

    filename = f"{tipo_artefato.upper()}_{projeto.id}_v{artefato.versao}.docx"
    
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@router.get("/{projeto_id}/{tipo_artefato}/print", response_class=HTMLResponse)
async def exportar_print(
    request: Request,
    projeto_id: int,
    tipo_artefato: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(auth_get_current_user)
):
    """Gera visualização de impressão (PDF via browser)"""
    if tipo_artefato not in ARTEFATO_MAP:
        raise HTTPException(status_code=404, detail="Tipo de artefato inválido")

    config = ARTEFATO_MAP[tipo_artefato]
    Model = config["model"]
    
    result = await db.execute(select(Projeto).filter(Projeto.id == projeto_id))
    projeto = result.scalars().first()

    result = await db.execute(select(Model).filter(Model.projeto_id == projeto_id))
    artefato = result.scalars().first()
    
    if not artefato:
        raise HTTPException(status_code=404, detail="Artefato não encontrado")

    return templates.TemplateResponse(
        "pages/artefato_print.html",
        {
            "request": request,
            "projeto": projeto,
            "artefato": artefato,
            "titulo": config["titulo"],
            "campos_config": config["config"],
            "data_impressao": datetime.now().strftime('%d/%m/%Y')
        }
    )


@router.get("/pdf/{tipo_artefato}/{artefato_id}")
async def exportar_pdf(
    tipo_artefato: str,
    artefato_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(auth_get_current_user)
):
    """
    Gera arquivo PDF do artefato de forma dinâmica (em thread separada).
    """
    if not HAS_WEASYPRINT:
        raise HTTPException(
            status_code=501, 
            detail="WeasyPrint não está instalado. Execute: pip install weasyprint"
        )
    
    if tipo_artefato not in ARTEFATO_MAP:
        raise HTTPException(
            status_code=400, 
            detail=f"Tipo de artefato inválido. Tipos válidos: {', '.join(ARTEFATO_MAP.keys())}"
        )
    
    config = ARTEFATO_MAP[tipo_artefato]
    Model = config["model"]
    campos_config = config["config"]
    titulo = config["titulo"]
    
    result = await db.execute(select(Model).filter(Model.id == artefato_id))
    artefato = result.scalars().first()
    if not artefato:
        raise HTTPException(
            status_code=404, 
            detail=f"Artefato {tipo_artefato} com ID {artefato_id} não encontrado"
        )
    
    result = await db.execute(select(Projeto).filter(Projeto.id == artefato.projeto_id))
    projeto = result.scalars().first()
    if not projeto:
        raise HTTPException(
            status_code=404, 
            detail="Projeto relacionado não encontrado"
        )
    
    try:
        # Executar geração de PDF (WebsyPrint é bloqueante) em thread separada
        pdf_buffer = await asyncio.to_thread(
            gerar_pdf_artefato,
            artefato=artefato,
            tipo_artefato=tipo_artefato,
            campos_config=campos_config,
            titulo=titulo,
            projeto=projeto
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Erro ao gerar PDF: {str(e)}"
        )
    
    filename = gerar_nome_arquivo_pdf(
        tipo_artefato=tipo_artefato,
        projeto_id=projeto.id,
        versao=artefato.versao
    )
    
    return StreamingResponse(
        pdf_buffer,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        }
    )